import os
from docx import Document
import requests
import pandas as pd


def fetch_player_data():
    # URL of the JSON data
    url = 'https://geo145327-staging.s3.ap-southeast-2.amazonaws.com/public/player_statistics_2024.json'

    # Send a GET request to the URL
    response = requests.get(url)

    # Check if the request was successful
    if response.status_code == 200:
        # Load the JSON data
        data = response.json()

        # Extract the first player's data
        first_player = data['PlayerStats'][0]['2024'][0]['0'][0]['2024-1-Sea-Eagles-v-Rabbitohs'][0]

        df = pd.DataFrame([first_player])
        return df

    else:
        print("Failed to retrieve data, status code:", response.status_code)
        return None


def create_player_tables(df):

    doc = Document()

    # Create a table with the player's basic info
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribute'
    hdr_cells[1].text = 'Value'

    # Add rows for basic player information
    for key in ['Name', 'Number', 'Position']:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    # Add a section for Per Game stats
    doc.add_paragraph("PER GAME STATS:")
    table_pg = doc.add_table(rows=1, cols=2)
    hdr_cells_pg = table_pg.rows[0].cells
    hdr_cells_pg[0].text = 'Statistic'
    hdr_cells_pg[1].text = 'Value'

    for key in ['Tries', 'Try Assists']:
        row_cells = table_pg.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    # Add a section for Total stats
    doc.add_paragraph("TOTAL STATS:")
    table_total = doc.add_table(rows=1, cols=2)
    hdr_cells_total = table_total.rows[0].cells
    hdr_cells_total[0].text = 'Statistic'
    hdr_cells_total[1].text = 'Value'

    for key in ['Total Points', 'All Run Metres', 'Offloads', 'Average Play The Ball Speed', 'Line Breaks', 'Passes', 'On Report']:
        row_cells = table_total.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    # Calculate averages and add them
    doc.add_paragraph("AVERAGES:")
    try:
        avg_tries_per_game = float(
            df.iloc[0]['Total Points']) / float(df.iloc[0]['Tries'])
        doc.add_paragraph(
            f"Average Tries Per Game: {avg_tries_per_game:.2f}")
    except (ValueError, ZeroDivisionError):
        doc.add_paragraph("Not enough data for average tries per game.")

    try:
        avg_points_per_metre = float(
            df.iloc[0]['Total Points']) / float(df.iloc[0]['All Run Metres'])
        doc.add_paragraph(
            f"Average Points Per Metre Ran: {avg_points_per_metre:.2f}")
    except (ValueError, ZeroDivisionError):
        doc.add_paragraph(
            "Not enough data for average points per metre ran.")

    return doc


df_player = fetch_player_data()

print(df_player)

if df_player is not None:
    new_doc = create_player_tables(df_player)
    filename = "df_player.docx"
    new_doc.save(filename)
    os.system(f"open '{filename}'")

# here im attempting to untie fetch process from create tables process so I can render data based on dataframes and not have them tied into the word render process

# then I can call the functions spearately within the generate_player function
