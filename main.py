import json
import warnings
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from match_scraper import get_match_data
from player_scraper import get_player_data

# Suppress specific deprecation warnings
warnings.filterwarnings('ignore', category=UserWarning,
                        message='.*style lookup by style_id is deprecated.*')


def create_word_document():
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('NRL Report', level=0)

    para = doc.add_paragraph()
    para.add_run('')

    # Add a centered image
    para_image = doc.add_paragraph()
    run_image = para_image.add_run()
    run_image.add_picture('nrl_image.jpg', width=Cm(12))
    para_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the paragraph

    # add a page
    doc.add_page_break()

    doc.add_paragraph('The History of NRL', style='Heading 1')

    # add content to the second page
    history_text = (
        "The National Rugby League (NRL) is the top league of professional rugby league men's clubs in Australasia. "
        "Run by the Australian Rugby League Commission, the NRL's main competition is known as the Telstra Premiership "
        "due to sponsorship from Telstra Corporation and is contested by sixteen teams, fifteen of which are based in "
        "Australia with one based in New Zealand. NRL games are played throughout Australia and New Zealand from March "
        "to October. The season culminates in the premiership-deciding game, the NRL Grand Final, traditionally one of "
        "Australia's most popular sporting events and one of the largest attended club championship events in the world."
    )

    doc.add_paragraph(history_text, style='BodyText')

    # Align paragraph
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # add a page
    doc.add_page_break()

    doc.add_paragraph('NRL MATCH DATA:', style='Heading 1')

    get_match_data(doc)

    # add a page
    doc.add_page_break()

    get_player_data(doc)

    filename = 'test.docx'

    # Save the document
    doc.save(filename)

    # Open the doc with the default application
    os.system(f"open '{filename}'")


# Run the function to create the document
create_word_document()
