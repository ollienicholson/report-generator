from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Averages Calculation and Display


def calculate_average(numerator, denominator):
    try:
        value = float(numerator) / float(denominator)
        return value
    except (ValueError, ZeroDivisionError):
        return None


def add_table(doc, column_titles=None, data_dict=None):
    """
    Helper function to add a table to the document. Row=1, cols=2.\n
    Function takes doc, column titles and a dict.\n
    Returns the doc object.
    """
    table = doc.add_table(rows=1 if column_titles else 0, cols=2)
    if column_titles:
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = column_titles[0]
        hdr_cells[1].text = column_titles[1]

    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
    return doc


def add_chart_to_doc(doc, filename):
    """
    Inserts a bar chart image into the Word document.
    """
    doc.add_paragraph()  # Add an empty paragraph for spacing
    doc.add_paragraph('Player Performance Stats:', style='Heading 2')
    para_image = doc.add_paragraph()
    run_image = para_image.add_run()
    run_image.add_picture(filename, width=Cm(12))
    para_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the paragraph


def on_enter(event):
    # Change background color on hover
    print("Mouse entered player button")  # Debug print
    event.widget.config(background='red')


def on_leave(event):
    # Change back to default on mouse leave
    print("Mouse left plyer button\n")  # Debug print
    event.widget.config(background='blue')
