import warnings
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from player_scraper import get_player_data

# Suppress specific deprecation warnings
warnings.filterwarnings('ignore', category=UserWarning,
                        message='.*style lookup by style_id is deprecated.*')


def create_player_report():
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('NRL Player Report', level=0)

    doc.add_paragraph()

    # Add a centered image
    para_image = doc.add_paragraph()
    run_image = para_image.add_run()
    run_image.add_picture('stock_image.jpg', width=Cm(12))
    para_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the paragraph

    # add a page
    doc.add_page_break()

    doc.add_paragraph('A little bit about the player', style='Heading 1')
    doc.add_paragraph()

    # add content to the second page
    history_text = (
        "Add a blurb about the player from ChatGPT"
    )

    doc.add_paragraph(history_text, style='BodyText')

    # Align paragraph
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # add a page
    doc.add_page_break()

    doc.add_paragraph('NRL PLAYER DATA:', style='Heading 1')

    get_player_data(doc)

    filename = 'player_report_test.docx'

    # Save the document
    doc.save(filename)

    # Open the doc with the default application
    os.system(f"open '{filename}'")