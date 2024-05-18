import warnings
import logging
import os

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from data.fetchers import fetch_player_data
from players.player_tables import create_player_tables
from testing.test_charts import player_stats_chart

from helpers.helpers import Helpers

# Suppress specific deprecation warnings
warnings.filterwarnings('ignore', category=UserWarning,
                        message='.*style lookup by style_id is deprecated.*')


def create_player_report():
    '''creates the player report

    TODO: pass dynamic images, pass dynamic player names based on GUI input

    '''
    logging.debug("Creating player report")

    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('NRL Player Report', level=0)

    doc.add_paragraph()

    # Add a centered image
    para_image = doc.add_paragraph()
    run_image = para_image.add_run()
    run_image.add_picture('images/stock_image.jpg', width=Cm(12))
    para_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the paragraph

    # add a page
    doc.add_page_break()

    doc.add_paragraph('A little bit about the player', style='Heading 1')
    doc.add_paragraph()

    # add content to the second page
    history_text = (
        "Add a blurb about the player from ChatGPT"
    )

    doc.add_paragraph(history_text, style='BodyText')

    # Align paragraph
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # add a page
    doc.add_page_break()

    doc.add_paragraph('NRL PLAYER DATA:', style='Heading 1')

    # fetch player data
    df_player = fetch_player_data()

    if df_player is not None:
        # insert player tables
        create_player_tables(df_player, doc)

    # Save player stats chart
        chart_filename = 'charts/player_stats.png'
        player_stats_chart(df_player, chart_filename)

        # Add the player stats chart to the document
        HELP = Helpers()
        HELP.add_player_chart(doc, chart_filename)

    # allocate save name, save location, save doc and open doc
    output_folder = 'report_outputs'
    filename = f"{output_folder}/player_report_test.docx"

    # Save the document
    doc.save(filename)

    # Open the doc with the default application
    os.system(f"open '{filename}'")
