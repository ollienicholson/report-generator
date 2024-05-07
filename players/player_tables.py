from docx import Document
import pandas as pd
from helpers import calculate_average, add_table


def create_player_tables(df: pd.DataFrame, doc: Document):
    """
    Function that accepts a document and a DataFrame object,
    creates data tables in the document, and returns the document.
    """

    # NOTE will need to review the line below when rendering specific player data
    player_info = df.iloc[0]  # Access the row once, and reuse it

    # Create a table with the player's basic info
    headers = None
    basic_info_keys = ['Name', 'Number', 'Position']
    basic_info = {key: player_info[key] for key in basic_info_keys}
    doc = add_table(doc, headers, basic_info)
    doc.add_paragraph()

    # PER GAME STATS
    doc.add_paragraph("PER GAME STATS:")
    game_stats_keys = ['Tries', 'Try Assists']
    game_stats = {key: player_info[key] for key in game_stats_keys}
    doc = add_table(doc, ['Statistic', 'Value'], game_stats)

    # TOTAL STATS
    doc.add_paragraph("TOTAL STATS:")
    total_stats_keys = ['Total Points', 'All Run Metres', 'Offloads',
                        'Average Play The Ball Speed', 'Line Breaks', 'Passes', 'On Report']
    total_stats = {key: player_info[key] for key in total_stats_keys}
    doc = add_table(doc, ['Statistic', 'Value'], total_stats)

    avg_tries_per_game = calculate_average(
        player_info['Total Points'], player_info['Tries'])
    avg_points_per_metre = calculate_average(
        player_info['Total Points'], player_info['All Run Metres'])

    doc.add_paragraph("AVERAGES:")
    if avg_tries_per_game is not None:
        doc.add_paragraph(f"Average Tries Per Game: {avg_tries_per_game:.2f}")
    else:
        doc.add_paragraph("Not enough data for average tries per game.")

    if avg_points_per_metre is not None:
        doc.add_paragraph(
            f"Average Points Per Metre Ran: {avg_points_per_metre:.2f}")
    else:
        doc.add_paragraph("Not enough data for average points per metre ran.")

    return doc
