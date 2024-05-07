from docx import Document
import pandas as pd
from helpers import calculate_average, add_table


def create_player_tables(df, doc):
    """
    function passes a doc & DF object,
    creates data tables in the doc
    returns the doc
    """
    # Create a table with the player's basic info
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribute'
    hdr_cells[1].text = 'Value'

    # Add rows for basic player information
    for key in ['Name', 'Number', 'Position']:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    # Add a section for Per Game stats
    doc.add_paragraph("PER GAME STATS:")
    table_pg = doc.add_table(rows=1, cols=2)
    hdr_cells_pg = table_pg.rows[0].cells
    hdr_cells_pg[0].text = 'Statistic'
    hdr_cells_pg[1].text = 'Value'

    for key in ['Tries', 'Try Assists']:
        row_cells = table_pg.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    # Add a section for Total stats
    doc.add_paragraph("TOTAL STATS:")
    table_total = doc.add_table(rows=1, cols=2)
    hdr_cells_total = table_total.rows[0].cells
    hdr_cells_total[0].text = 'Statistic'
    hdr_cells_total[1].text = 'Value'

    for key in ['Total Points', 'All Run Metres', 'Offloads', 'Average Play The Ball Speed', 'Line Breaks', 'Passes', 'On Report']:
        row_cells = table_total.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(df.iloc[0][key])

    avg_tries_per_game = calculate_average(
        df.iloc[0]['Total Points'], df.iloc[0]['Tries'])
    avg_points_per_metre = calculate_average(
        df.iloc[0]['Total Points'], df.iloc[0]['All Run Metres'])

    # Calculate averages and add them
    doc.add_paragraph("AVERAGES:")

    if avg_tries_per_game is not None:
        doc.add_paragraph(f"Average Tries Per Game: {avg_tries_per_game:.2f}")
    else:
        doc.add_paragraph("Not enough data for average tries per game.")

    if avg_points_per_metre is not None:
        doc.add_paragraph(
            f"Average Points Per Metre Ran: {avg_points_per_metre:.2f}")
    else:
        doc.add_paragraph("Not enough data for average points per metre ran.")

    return doc


# df_player = fetch_player_data()

# if df_player is not None:
#     doc = Document()
#     new_doc = create_player_tables(df_player, doc)
#     filename = "df_player.docx"
#     new_doc.save(filename)
#     os.system(f"open '{filename}'")

# TEST FOR DECOUPLING PLAYER SCRAPER
#
# here im attempting to untie fetch process from create tables process so I can render data based on dataframes and not have them tied into the word render process

# then I can call the functions spearately within the generate_player function

# def create_player_tables(df, doc):
#     """
#     Function that accepts a document and a DataFrame object,
#     creates data tables in the document, and returns the document.
#     """
#     player_info = df.iloc[0]  # Access the row once, and reuse it

#     # Create a table with the player's basic info
#     basic_info_keys = ['Name', 'Number', 'Position']
#     basic_info = {key: player_info[key] for key in basic_info_keys}
#     doc = add_table(doc, ['Attribute', 'Value'], basic_info)

#     # PER GAME STATS
#     doc.add_paragraph("PER GAME STATS:")
#     game_stats_keys = ['Tries', 'Try Assists']
#     game_stats = {key: player_info[key] for key in game_stats_keys}
#     doc = add_table(doc, ['Statistic', 'Value'], game_stats)

#     # TOTAL STATS
#     doc.add_paragraph("TOTAL STATS:")
#     total_stats_keys = ['Total Points', 'All Run Metres', 'Offloads',
#                         'Average Play The Ball Speed', 'Line Breaks', 'Passes', 'On Report']
#     total_stats = {key: player_info[key] for key in total_stats_keys}
#     doc = add_table(doc, ['Statistic', 'Value'], total_stats)

#     # Averages Calculation and Display
#     def calculate_average(numerator, denominator):
#         try:
#             return float(numerator) / float(denominator)
#         except (ValueError, ZeroDivisionError):
#             return None

#     avg_tries_per_game = calculate_average(
#         player_info['Total Points'], player_info['Tries'])
#     avg_points_per_metre = calculate_average(
#         player_info['Total Points'], player_info['All Run Metres'])

#     doc.add_paragraph("AVERAGES:")
#     if avg_tries_per_game is not None:
#         doc.add_paragraph(f"Average Tries Per Game: {avg_tries_per_game:.2f}")
#     else:
#         doc.add_paragraph("Not enough data for average tries per game.")

#     if avg_points_per_metre is not None:
#         doc.add_paragraph(
#             f"Average Points Per Metre Ran: {avg_points_per_metre:.2f}")
#     else:
#         doc.add_paragraph("Not enough data for average points per metre ran.")

#     return doc
